"""File service: docx upload/edit + image upload for vision.

Uses python-docx to read and modify .docx files while preserving formatting.
Modified files are saved as NEW files (the original is never touched).
Images are saved as-is and served via a view endpoint for display.
"""

from __future__ import annotations

import base64
import mimetypes
import uuid
from pathlib import Path

from sqlalchemy.orm import Session

from app.config import settings
from app.models.database import File

UPLOAD_DIR = settings.data_dir / "uploads"
UPLOAD_DIR.mkdir(parents=True, exist_ok=True)

# Supported image formats
IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".gif", ".webp", ".bmp"}
MAX_IMAGE_SIZE = 10 * 1024 * 1024  # 10MB


class FileError(Exception):
    """Raised when a file operation fails."""


def _new_id() -> str:
    return uuid.uuid4().hex


# --------------------------------------------------------------------------- #
#  File records (DB)
# --------------------------------------------------------------------------- #

def create_file_record(
    db: Session,
    *,
    filename: str,
    path: str,
    role: str,
    original_id: str | None = None,
) -> File:
    rec = File(id=_new_id(), filename=filename, path=path, role=role, original_id=original_id)
    db.add(rec)
    db.commit()
    db.refresh(rec)
    return rec


def get_file_record(db: Session, file_id: str) -> File | None:
    return db.get(File, file_id)


# --------------------------------------------------------------------------- #
#  Upload
# --------------------------------------------------------------------------- #

def save_docx_upload(content: bytes, filename: str) -> dict:
    """Save an uploaded docx; returns {filename, path, role}."""
    if not filename.lower().endswith(".docx"):
        raise FileError("仅支持 .docx 文件")
    safe_name = Path(filename).name  # strip any directory components
    fid = _new_id()
    path = UPLOAD_DIR / f"{fid}.docx"
    path.write_bytes(content)
    return {"filename": safe_name, "path": str(path), "role": "upload"}


# --------------------------------------------------------------------------- #
#  Extraction
# --------------------------------------------------------------------------- #

def extract_docx_paragraphs(path: str) -> list[str]:
    """Extract paragraph texts from a docx (non-empty)."""
    from docx import Document
    doc = Document(path)
    return [p.text.strip() for p in doc.paragraphs if p.text.strip()]


def paragraph_indexed_text(path: str, max_paras: int = 200, max_chars: int = 12000) -> str:
    """Build a numbered paragraph listing for the LLM context."""
    paras = extract_docx_paragraphs(path)
    lines: list[str] = []
    total = 0
    for i, t in enumerate(paras):
        if i >= max_paras:
            lines.append("……（段落较多，已截断）")
            break
        lines.append(f"[{i}] {t}")
        total += len(t)
        if total > max_chars:
            lines.append("……（文档较长，已截断）")
            break
    return "\n".join(lines) if lines else "（文档为空）"


# --------------------------------------------------------------------------- #
#  Editing (creates a NEW modified file)
# --------------------------------------------------------------------------- #

def apply_docx_edit(src_path: str, paragraph_index: int, new_text: str, filename: str) -> dict:
    """Replace one paragraph's text in a copy of the docx.

    Returns a new file record payload {filename, path, role}.
    """
    from docx import Document

    if paragraph_index < 0:
        raise FileError("段落索引无效")
    doc = Document(src_path)
    paras = doc.paragraphs
    if paragraph_index >= len(paras):
        raise FileError(f"段落索引超出范围（共 {len(paras)} 段，索引从 0 开始）")

    target = paras[paragraph_index]
    if target.runs:
        # Keep the first run's formatting; apply new text to it and clear the rest
        target.runs[0].text = new_text
        for r in target.runs[1:]:
            r.text = ""
    else:
        target.add_run(new_text)

    fid = _new_id()
    path = UPLOAD_DIR / f"{fid}.docx"
    doc.save(str(path))

    stem = Path(filename).stem or "document"
    ext = Path(filename).suffix or ".docx"
    new_name = f"{stem}_modified{ext}"
    return {"filename": new_name, "path": str(path), "role": "generated"}


# --------------------------------------------------------------------------- #
#  Image upload (for vision / OCR)
# --------------------------------------------------------------------------- #

def save_image_upload(content: bytes, filename: str) -> dict:
    """Save an uploaded image; returns {filename, path, role}.

    Validates extension and size. The original extension is preserved.
    """
    ext = Path(filename).suffix.lower()
    if ext not in IMAGE_EXTENSIONS:
        raise FileError(f"不支持的图片格式 {ext}，支持：{', '.join(sorted(IMAGE_EXTENSIONS))}")
    if len(content) > MAX_IMAGE_SIZE:
        raise FileError(f"图片过大（{len(content) // 1024 // 1024}MB），最大支持 {MAX_IMAGE_SIZE // 1024 // 1024}MB")

    safe_name = Path(filename).name
    fid = _new_id()
    path = UPLOAD_DIR / f"{fid}{ext}"
    path.write_bytes(content)
    return {"filename": safe_name, "path": str(path), "role": "upload"}


def image_to_base64_url(path: str) -> str:
    """Read an image file and return a base64 data URL for the LLM API."""
    p = Path(path)
    if not p.exists():
        raise FileError("图片文件不存在")
    content = p.read_bytes()
    ext = p.suffix.lower()
    mime = mimetypes.types_map.get(ext, "image/jpeg")
    b64 = base64.b64encode(content).decode("ascii")
    return f"data:{mime};base64,{b64}"